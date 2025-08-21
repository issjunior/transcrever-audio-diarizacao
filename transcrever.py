import os
import warnings
from dotenv import load_dotenv
import whisper
from pyannote.audio import Pipeline
from docx import Document
from tqdm import tqdm  # barra de progresso
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------------------
# 0. Supressão de warnings
# -------------------------------
warnings.filterwarnings("ignore")  # ignora warnings de torchaudio, PyTorch, etc.

# -------------------------------
# 1. Carregar variáveis de ambiente
# -------------------------------
load_dotenv()  # Carrega as variáveis do arquivo .env
HUGGINGFACE_TOKEN = os.getenv("HUGGINGFACE_TOKEN")

if not HUGGINGFACE_TOKEN:
    raise ValueError("❌ Token do HuggingFace não encontrado. Defina no arquivo .env")

# -------------------------------
# 2. Transcrição com Whisper
# -------------------------------
modelo = whisper.load_model("large")  # tiny, base, small, medium, large
audio_path = "exemplo_pequeno.mp3"  # Pyannote/torchaudio somente aceita arquivos de áudio

print(">> Rodando Whisper...")
resultado = modelo.transcribe(audio_path)

# -------------------------------
# 3. Diarização com Pyannote
# -------------------------------
print(">> Rodando diarização...")

pipeline = Pipeline.from_pretrained(
    "pyannote/speaker-diarization",
    use_auth_token=HUGGINGFACE_TOKEN
)

diarization = pipeline(audio_path)

# -------------------------------
# Função para formatar tempo em mm:ss
# -------------------------------
def formatar_tempo(segundos):
    minutos = int(segundos // 60)
    segundos_restantes = int(segundos % 60)
    return f"{minutos:02d}:{segundos_restantes:02d}"

# -------------------------------
# 4. Mesclar transcrição + locutores numerados
# -------------------------------
falas = []

# Dicionário para mapear locutores originais para "Locutor 1", "Locutor 2", ...
mapa_locutores = {}
contador_locutor = 1

print(">> Mesclando transcrição e locutores...")
for segmento in tqdm(resultado["segments"], desc="Processando segmentos"):
    start = segmento["start"]
    end = segmento["end"]
    texto = segmento["text"].strip()

    speaker = "Desconhecido"
    for turno in diarization.itertracks(yield_label=True):
        seg_dia = turno[0]        # primeiro elemento é sempre o segmento
        locutor_original = turno[-1]       # último elemento é o label/locutor
        s = seg_dia.start
        e = seg_dia.end
        if s <= start <= e:
            speaker = locutor_original
            break

    # Mapear locutor para "Locutor X"
    if speaker not in mapa_locutores and speaker != "Desconhecido":
        mapa_locutores[speaker] = f"Locutor {contador_locutor}"
        contador_locutor += 1

    nome_final = mapa_locutores.get(speaker, speaker)

    falas.append({
        "tempo": f"{formatar_tempo(start)} - {formatar_tempo(end)}",
        "locutor": nome_final,
        "texto": texto
    })

# Função corrigida para adicionar bordas à tabela
def adicionar_bordas_tabela(tabela):
    tbl = tabela._element
    tbl_pr = tbl.find(qn('w:tblPr'))  # Tenta encontrar o elemento tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')  # Cria tblPr se não existir
        tbl.insert(0, tbl_pr)

    tbl_borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # Tipo de borda: 'single' (linha simples)
        border.set(qn('w:sz'), '4')       # Tamanho da borda
        border.set(qn('w:space'), '0')    # Espaçamento
        border.set(qn('w:color'), '000000')  # Cor: preto
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)

# -------------------------------
# 5. Exportar para Word
# -------------------------------
print(">> Exportando para Word...")
doc = Document()
doc.add_heading("Tabela 1 - transcrição de áudio", level=1)

tabela = doc.add_table(rows=1, cols=3)
hdr_cells = tabela.rows[0].cells
hdr_cells[0].text = 'Tempo'
hdr_cells[1].text = 'Locutor'
hdr_cells[2].text = 'Transcrição'

for fala in falas:
    row_cells = tabela.add_row().cells
    row_cells[0].text = fala["tempo"]
    row_cells[1].text = fala["locutor"]
    row_cells[2].text = fala["texto"]

# Adicionar bordas à tabela
adicionar_bordas_tabela(tabela)

doc.save("transcricao_diarizada.docx")
print("✅ Arquivo gerado: transcricao_diarizada.docx")
